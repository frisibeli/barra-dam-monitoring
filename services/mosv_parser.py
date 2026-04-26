import re
import subprocess
from pathlib import Path

# Matches the Огоста data line in the standardised bulletin table.
# Uses \s* (not \s+) because .doc files extracted via textutil may have
# no whitespace between columns.
#
# Number format in bulletins: "506,000" or "37,01" — always comma as
# decimal separator, 1-3 integer digits, comma, 1-3 fractional digits.
_NUM = r'(\d{1,3},\d{1,3})'
_PCT = r'(\d{1,3},\d{1,3})%'
_NUM_SRC = r'\d{1,3},\d{1,3}'

_OGOSTA_LINE_RE = re.compile(
    r'Огоста\s*'
    + _NUM + r'\s*'   # total_capacity   506,000
    + _NUM + r'\s*'   # dead_volume       67,000
    + _NUM + r'\s*'   # current_volume   166,571
    + _PCT + r'\s*'   # pct_total         32,92%
    + _NUM + r'\s*'   # useful_volume     99,571
    + _PCT + r'\s*'   # pct_useful        22,68%
    + _NUM + r'\s*'   # inflow             6,586
    + _NUM,           # outflow            3,542
)

# ---------------------------------------------------------------------------
# All-dams registry
# Maps the Bulgarian name as it appears in the bulletin → stable slug.
# Names containing partial overlap (e.g. "Искър" / "Бели Искър") are handled
# by the per-dam regex using a Cyrillic lookbehind so "Искър" won't match
# inside "Бели Искър", and a negative lookahead so "Белмекен" won't match
# inside "Белмекен-Чаира".
# ---------------------------------------------------------------------------
_DAM_REGISTRY: dict[str, str] = {
    "Искър": "iskar",
    "Бели Искър": "beli_iskar",
    "Среченска бара": "srechenska_bara",
    "Христо Смирненски": "hristo_smirnenski",
    "Йовковци": "yovkovtsi",
    "Тича": "ticha",
    "Камчия": "kamchia",
    "Ясна поляна": "yasna_polyana",
    "Асеновец": "asenovets",
    "Боровица": "borovitsa",
    "Студена": "studena",
    "Дяково": "dyakovo",
    "Калин": "kalin",
    "Карагьол": "karagyol",
    "Огняново": "ognyanova",
    "Порой": "poroy",
    "Ахелой": "akheloy",
    "Панчарево": "pancharevo",
    "Ястребино": "yastrebino",
    "Кула": "kula",
    "Рабиша": "rabisha",
    "Огоста": "ogosta",
    "Сопот": "sopot",
    "Горни Дъбник": "gorni_dabnik",
    "Бели Лом": "beli_lom",
    "Съединение": "saedenie",
    "Георги Трайков": "georgi_traikov",
    "Жребчево": "zhrebchevo",
    "Малко Шарково": "malko_sharkovo",
    "Домлян": "domlyan",
    "Пясъчник": "pyasachnik",
    "Тополница": "topolnitsa",
    "Тракиец": "trakiets",
    "Пчелина": "pchelina",
    "Александър Стамболийски": "aleksandar_stamboliyski",
    "Кокаляне": "kokalyane",
    "Копринка": "koprinka",
    "Белмекен": "belmeken",
    "Чаира": "chaira",
    "Голям Беглик": "golyam_beglik",
    "Широка поляна": "shiroka_polyana",
    "Беглика": "beglika",
    "Тошков Чарк": "toshkov_chark",
    "Батак": "batak",
    "Доспат": "dospat",
    "Цанков камък": "tsankov_kamak",
    "Въча": "vacha",
    "Кричим": "krichim",
    "Кърджали": "kardzhali",
    "Студен кладенец": "studen_kladenets",
    "Ивайловград": "ivaylovgrad",
    "Розов кладенец": "rozov_kladenets",
}


def _build_dam_re(name_bg: str) -> re.Pattern:
    """Build a regex that matches a dam data row for the given Bulgarian name.

    Uses:
    - Negative lookbehind (?<![а-яА-Я]) so "Искър" doesn't match inside "Бели Искър".
    - Negative lookahead (?![а-яА-Я\\-]) so "Белмекен" doesn't match inside
      "Белмекен-Чаира", and "Голям Беглик" doesn't match inside
      "Голям Беглик-Широка поляна".
    - [^0-9\\n]{0,80} absorbs optional suffix text before the first number
      (e.g. "- за рез.водоснабдяване **").
    - Inflow and outflow are optional — some dams report only 6 columns.
    """
    return re.compile(
        r'(?<![а-яА-ЯёЁ\-])' + re.escape(name_bg) + r'(?![а-яА-ЯёЁ\-])'
        + r'[^0-9\n]{0,80}'                 # optional suffix before first digit
        + r'(' + _NUM_SRC + r')\s+'         # total_capacity
        + r'(' + _NUM_SRC + r')\s+'         # dead_volume
        + r'(' + _NUM_SRC + r')\s+'         # current_volume (Наличен)
        + r'(' + _NUM_SRC + r')%\s+'        # pct_total
        + r'(' + _NUM_SRC + r')\s+'         # useful_volume (Разполагаем)
        + r'(' + _NUM_SRC + r')%'           # pct_useful
        + r'(?:\s+(' + _NUM_SRC + r')\s+(' + _NUM_SRC + r'))?',  # optional inflow+outflow
    )


# Pre-compiled per-dam regexes (built once at import time)
_DAM_RES: dict[str, re.Pattern] = {
    name: _build_dam_re(name) for name in _DAM_REGISTRY
}


def _bg_float(s: str) -> float:
    return float(s.replace(",", "."))


class MosvParserService:
    def extract_volume(self, path: Path) -> dict | None:
        """Try all extraction strategies in order of reliability."""
        try:
            text = self._extract_text(path)
            result = self._extract_from_line(text)
            if result:
                return result
        except Exception as e:
            print(f"  Text extraction failed: {e}")

        if path.suffix.lower() == ".docx":
            try:
                return self._extract_from_docx_tables(path)
            except Exception as e:
                print(f"  Table extraction failed: {e}")

        return None

    def extract_all_volumes(self, path: Path) -> dict[str, dict]:
        """Extract volume data for all known dams from a bulletin file.

        Returns a dict mapping slug → reading-fields dict (same structure as
        extract_volume but with the dam slug as key).  Dams for which data
        cannot be found (e.g. Кокаляне when it has no data) are omitted.
        """
        try:
            text = self._extract_text(path)
        except Exception as e:
            print(f"  Text extraction failed: {e}")
            return {}

        results: dict[str, dict] = {}
        for name_bg, slug in _DAM_REGISTRY.items():
            data = self._extract_dam_from_text(text, name_bg)
            if data:
                results[slug] = data

        # Docx table fallback: pick up anything missed
        if path.suffix.lower() == ".docx" and not results:
            try:
                results = self._extract_all_from_docx_tables(path)
            except Exception as e:
                print(f"  Table extraction failed: {e}")

        return results

    def _extract_dam_from_text(self, text: str, name_bg: str) -> dict | None:
        """Parse one dam's row from extracted bulletin text."""
        pattern = _DAM_RES.get(name_bg)
        if pattern is None:
            return None
        m = pattern.search(text)
        if not m:
            return None
        g = m.groups()
        # g[0..5] are the 6 mandatory columns; g[6,7] are optional inflow/outflow
        result: dict = {
            "total_capacity_mm3": _bg_float(g[0]),
            "dead_volume_mm3": _bg_float(g[1]),
            "volume_mm3": _bg_float(g[2]),
            "pct_total": _bg_float(g[3]),
            "useful_volume_mm3": _bg_float(g[4]),
            "pct_useful": _bg_float(g[5]),
        }
        if g[6] is not None:
            result["inflow_m3s"] = _bg_float(g[6])
        if g[7] is not None:
            result["outflow_m3s"] = _bg_float(g[7])
        return result

    def _extract_all_from_docx_tables(self, path: Path) -> dict[str, dict]:
        """Docx table fallback: scan all rows and try to match any registered dam."""
        if path.suffix.lower() != ".docx":
            return {}
        from docx import Document
        doc = Document(str(path))
        results: dict[str, dict] = {}
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                joined = " ".join(cells)
                for name_bg, slug in _DAM_REGISTRY.items():
                    if name_bg in joined and slug not in results:
                        data = self._extract_dam_from_text(joined, name_bg)
                        if data:
                            results[slug] = data
        return results

    def _extract_text(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext == ".docx":
            return self._text_from_docx(path)
        elif ext == ".doc":
            return self._text_from_doc(path)
        elif ext == ".pdf":
            return self._text_from_pdf(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _text_from_docx(self, path: Path) -> str:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    def _text_from_doc(self, path: Path) -> str:
        # macOS built-in textutil converts legacy .doc to plain text
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True,
            timeout=30,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"textutil failed (exit {result.returncode}): {result.stderr.decode()}"
            )
        text = result.stdout.decode("utf-8", errors="replace")
        return text.replace("\x07", " ")  # textutil uses BEL as cell separator

    def _text_from_pdf(self, path: Path) -> str:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    def _extract_from_line(self, text: str) -> dict | None:
        m = _OGOSTA_LINE_RE.search(text)
        if not m:
            return None
        return {
            "volume_mm3": _bg_float(m.group(3)),
            "total_capacity_mm3": _bg_float(m.group(1)),
            "dead_volume_mm3": _bg_float(m.group(2)),
            "pct_total": _bg_float(m.group(4)),
            "useful_volume_mm3": _bg_float(m.group(5)),
            "pct_useful": _bg_float(m.group(6)),
            "inflow_m3s": _bg_float(m.group(7)),
            "outflow_m3s": _bg_float(m.group(8)),
        }

    def _extract_from_docx_tables(self, path: Path) -> dict | None:
        if path.suffix.lower() != ".docx":
            return None
        from docx import Document
        doc = Document(str(path))
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                joined = " ".join(cells)
                if "Огоста" not in joined:
                    continue
                result = self._extract_from_line(joined)
                if result:
                    return result
                for cell in cells:
                    nums = re.findall(r'\d[\d\s]*[,\.]\d+', cell)
                    for raw in nums:
                        try:
                            vol = _bg_float(raw.replace(" ", ""))
                            if 0 < vol < 600:
                                return {"volume_mm3": vol}
                        except ValueError:
                            continue
        return None
